import json
import re
import numpy as np
import pandas as pd
from typing import List, Dict, Any, Tuple, Optional
from dataclasses import dataclass
from datetime import datetime, timedelta
import pickle
import logging
from pathlib import Path

# External libraries (install via pip)
# Initialize variables to None to avoid NameError
torch = None
AutoTokenizer = None
AutoModel = None
pipeline = None
SentenceTransformer = None
faiss = None
cosine_similarity = None
spacy = None
RecursiveCharacterTextSplitter = None
Document = None
PyPDF2 = None
DocxDocument = None
email = None
policy = None
fitz = None
transformers = None

try:
    import torch
    from transformers import AutoTokenizer, AutoModel, pipeline
    from sentence_transformers import SentenceTransformer
    import faiss
    from sklearn.metrics.pairwise import cosine_similarity
    import spacy
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.docstore.document import Document
    import PyPDF2
    from docx import Document as DocxDocument
    import email
    from email import policy
    import fitz  # PyMuPDF for better PDF processing
    import transformers
except ImportError as e:
    print(f"Missing dependencies. Install with: pip install {e.name}")
    print("Some features may not work without these dependencies.")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class QueryStructured:
    """Structured representation of parsed query"""
    age: Optional[int] = None
    gender: Optional[str] = None
    procedure: Optional[str] = None
    location: Optional[str] = None
    policy_duration: Optional[str] = None
    policy_age_months: Optional[int] = None
    raw_query: str = ""
    extracted_entities: Dict[str, Any] = None

@dataclass
class RetrievedClause:
    """Retrieved document clause with metadata"""
    content: str
    source_document: str
    page_number: Optional[int] = None
    section: Optional[str] = None
    relevance_score: float = 0.0
    clause_type: Optional[str] = None
    semantic_similarity: float = 0.0
    keyword_matches: List[str] = None

@dataclass
class DecisionResponse:
    """Final decision response structure"""
    decision: str  # "approved", "rejected", "requires_review"
    amount: Optional[float] = None
    confidence: float = 0.0
    justification: str = ""
    relevant_clauses: List[RetrievedClause] = None
    reasoning_chain: List[str] = None
    coverage_details: Dict[str, Any] = None

class EnhancedQueryParser:
    """Advanced NLP-based query parser using multiple techniques"""
    
    def __init__(self):
        # Load spaCy model for NER
        self.nlp = None
        if spacy is not None:
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except OSError:
                logger.warning("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            except Exception as e:
                logger.warning(f"spaCy not available: {e}")
        else:
            logger.warning("spaCy not installed. Install with: pip install spacy")
        
        # Initialize medical NER pipeline
        self.medical_ner = None
        if pipeline is not None:
            try:
                self.medical_ner = pipeline("ner", 
                                          model="d4data/biomedical-ner-all",
                                          aggregation_strategy="simple")
            except Exception as e:
                logger.warning(f"Medical NER model not available: {e}")
        else:
            logger.warning("Transformers pipeline not available. Install with: pip install transformers")
        
        # Enhanced medical terminology dictionary
        self.medical_terms = {
            'surgery': ['surgery', 'operation', 'procedure', 'surgical'],
            'knee': ['knee', 'knee surgery', 'arthroscopy', 'knee replacement', 'knee arthroplasty'],
            'cardiac': ['heart', 'cardiac', 'bypass', 'angioplasty', 'heart surgery'],
            'dental': ['dental', 'tooth', 'root canal', 'extraction', 'dental surgery'],
            'orthopedic': ['orthopedic', 'bone', 'fracture', 'joint', 'orthopedic surgery'],
            'neurological': ['brain', 'neurological', 'stroke', 'seizure', 'brain surgery']
        }
        
        # Enhanced location patterns for Indian cities
        self.indian_cities = [
            'pune', 'mumbai', 'delhi', 'bangalore', 'chennai', 'hyderabad', 
            'kolkata', 'ahmedabad', 'nagpur', 'indore', 'bhopal', 'lucknow',
            'kanpur', 'patna', 'jaipur', 'chandigarh', 'gurgaon', 'noida'
        ]
    
    def parse_query(self, query: str) -> QueryStructured:
        """Parse natural language query into structured format with enhanced extraction"""
        structured = QueryStructured(raw_query=query)
        
        # Enhanced regex patterns for extracting information
        # Handle abbreviated formats like "46M, knee surgery, Pune, 3-month policy"
        age_pattern = r'(\d+)[-\s]*(?:year|yr|y)?\s*(?:old)?(?:\s*(?:male|female|M|F))?'
        gender_pattern = r'\b(?:male|female|M|F|man|woman)\b'
        location_pattern = r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b'
        policy_duration_pattern = r'(\d+)[-\s]*(?:month|mon|m|year|yr|y)[-\s]*(?:old)?\s*(?:policy|insurance)'
        
        # Enhanced age extraction - handle formats like "46M", "46-year-old", "46 years"
        age_match = re.search(age_pattern, query, re.IGNORECASE)
        if age_match:
            structured.age = int(age_match.group(1))
        
        # Enhanced gender extraction - handle single letter formats
        gender_match = re.search(gender_pattern, query, re.IGNORECASE)
        if gender_match:
            gender = gender_match.group().lower()
            structured.gender = 'male' if gender in ['male', 'm', 'man'] else 'female'
        else:
            # Check for single letter gender indicators after age
            age_gender_pattern = r'(\d+)(M|F)\b'
            age_gender_match = re.search(age_gender_pattern, query, re.IGNORECASE)
            if age_gender_match:
                structured.age = int(age_gender_match.group(1))
                structured.gender = 'male' if age_gender_match.group(2).upper() == 'M' else 'female'
        
        # Enhanced policy duration extraction
        policy_match = re.search(policy_duration_pattern, query, re.IGNORECASE)
        if policy_match:
            duration = int(policy_match.group(1))
            structured.policy_age_months = duration
            structured.policy_duration = f"{duration} months"
        else:
            # Handle abbreviated formats like "3-month policy"
            short_policy_pattern = r'(\d+)[-\s]*(?:month|mon|m)[-\s]*(?:policy|insurance)'
            short_match = re.search(short_policy_pattern, query, re.IGNORECASE)
            if short_match:
                duration = int(short_match.group(1))
                structured.policy_age_months = duration
                structured.policy_duration = f"{duration} months"
        
        # Enhanced procedure extraction using medical terminology
        structured.procedure = self._extract_medical_procedure(query)
        
        # Enhanced location extraction with Indian cities
        structured.location = self._extract_location(query)
        
        # Use spaCy for advanced entity extraction
        if self.nlp:
            doc = self.nlp(query)
            entities = {}
            for ent in doc.ents:
                entities[ent.label_] = ent.text
            structured.extracted_entities = entities
        
        return structured
    
    def _extract_medical_procedure(self, query: str) -> Optional[str]:
        """Enhanced medical procedure extraction"""
        query_lower = query.lower()
        
        # Use medical NER if available
        if self.medical_ner:
            try:
                medical_entities = self.medical_ner(query)
                for entity in medical_entities:
                    if entity['entity_group'] in ['TREATMENT', 'PROCEDURE']:
                        return entity['word']
            except Exception as e:
                logger.warning(f"Medical NER failed: {e}")
        
        # Enhanced pattern matching with medical terminology
        for category, terms in self.medical_terms.items():
            for term in terms:
                if term in query_lower:
                    # Extract surrounding context for better procedure identification
                    words = query_lower.split()
                    try:
                        idx = words.index(term)
                        if idx > 0:
                            # Try to get more context (2 words before)
                            if idx > 1:
                                return f"{words[idx-2]} {words[idx-1]} {term}"
                            else:
                                return f"{words[idx-1]} {term}"
                        else:
                            return term
                    except ValueError:
                        # If exact word not found, use the term as is
                        return term
        
        # Fallback to basic medical term detection
        medical_terms = ['surgery', 'operation', 'procedure', 'treatment', 'therapy']
        for term in medical_terms:
            if term in query_lower:
                return term
        
        return None
    
    def _extract_location(self, query: str) -> Optional[str]:
        """Enhanced location extraction with Indian cities"""
        query_lower = query.lower()
        
        # Check for Indian cities first
        for city in self.indian_cities:
            if city in query_lower:
                return city.title()
        
        # Fallback to general location extraction
        words = query.split()
        potential_locations = []
        
        for word in words:
            if word[0].isupper() and len(word) > 2 and not any(char.isdigit() for char in word):
                potential_locations.append(word)
        
        # Check for multi-word locations
        for i in range(len(words) - 1):
            if (words[i][0].isupper() and words[i+1][0].isupper() and 
                len(words[i]) > 2 and len(words[i+1]) > 2):
                potential_locations.append(f"{words[i]} {words[i+1]}")
        
        if potential_locations:
            return potential_locations[0]  # Take first capitalized word
        
        return None

class MultiModelEmbeddingRetriever:
    """Advanced retrieval system using multiple embedding models and hybrid search"""
    
    def __init__(self, embedding_models: List[str] = None):
        if embedding_models is None:
            # Use multiple models for better accuracy
            embedding_models = [
                'all-MiniLM-L6-v2',  # Fast and efficient
                'all-mpnet-base-v2',  # High quality
                'paraphrase-multilingual-MiniLM-L12-v2'  # Multilingual support
            ]
        
        self.embedding_models = {}
        self.vector_stores = {}
        
        # Load embedding models with memory optimization
        if SentenceTransformer is not None:
            for model_name in embedding_models:
                try:
                    # Use device='cpu' to avoid GPU memory issues
                    self.embedding_models[model_name] = SentenceTransformer(model_name, device='cpu')
                    logger.info(f"Loaded embedding model: {model_name}")
                except Exception as e:
                    logger.warning(f"Failed to load {model_name}: {e}")
        else:
            logger.warning("SentenceTransformer not available. Install with: pip install sentence-transformers")
        
        # Enhanced text splitter for better chunking
        self.text_splitter = None
        if RecursiveCharacterTextSplitter is not None:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=512,
                chunk_overlap=100,  # Increased overlap for better context
                separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
            )
        else:
            logger.warning("RecursiveCharacterTextSplitter not available. Install with: pip install langchain")
        
        self.documents = []
        self.chunks = []
        self.chunk_embeddings = {}
        
        # Enhanced keyword indexing for hybrid search
        self.keyword_index = {}
        self.coverage_keywords = [
            'covered', 'eligible', 'approved', 'excluded', 'not covered',
            'waiting period', 'pre-existing', 'maximum', 'limit', 'coverage',
            'policy', 'insurance', 'claim', 'benefit', 'deductible'
        ]
    
    def load_documents(self, document_paths: List[str]):
        """Load and process documents from various formats with enhanced processing"""
        if self.text_splitter is None:
            logger.error("Text splitter not available. Cannot load documents.")
            return
            
        for path in document_paths:
            path_obj = Path(path)
            
            if path_obj.suffix.lower() == '.pdf':
                # Enhanced PDF processing for large documents
                self._extract_pdf_content_enhanced(path)
            elif path_obj.suffix.lower() in ['.docx', '.doc']:
                content = self._extract_docx_content(path)
                self._process_document_content(content, path)
            elif path_obj.suffix.lower() == '.txt':
                with open(path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self._process_document_content(content, path)
            else:
                logger.warning(f"Unsupported file format: {path}")
                continue
        
        logger.info(f"Loaded {len(self.documents)} documents, {len(self.chunks)} chunks")
        
        # Create embeddings for all chunks
        self._create_embeddings()
        
        # Build keyword index for hybrid search
        self._build_keyword_index()
    
    def _process_document_content(self, content: str, path: str):
        """Process document content and create chunks with enhanced metadata"""
        if Document is not None:
            doc = Document(page_content=content, metadata={'source': path})
            self.documents.append(doc)
            
            # Split into chunks
            chunks = self.text_splitter.split_documents([doc])
            
            # Add enhanced metadata to chunks
            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    'chunk_id': i,
                    'total_chunks': len(chunks),
                    'content_length': len(chunk.page_content)
                })
            
            self.chunks.extend(chunks)
        else:
            logger.error("Document class not available. Cannot process documents.")
    
    def _extract_pdf_content_enhanced(self, path: str):
        """Enhanced PDF extraction that processes all pages with better structure"""
        if fitz is None:
            logger.error("PyMuPDF not available. Install with: pip install PyMuPDF")
            return
            
        try:
            doc = fitz.open(path)
            total_pages = len(doc)
            
            logger.info(f"Processing PDF with {total_pages} pages...")
            
            # Process each page individually for better chunking
            for page_num in range(total_pages):
                page = doc[page_num]
                text = page.get_text()
                
                if text.strip():  # Only process non-empty pages
                    # Create document for this page
                    if Document is not None:
                        page_doc = Document(
                            page_content=text, 
                            metadata={
                                'source': path,
                                'page_number': page_num + 1,
                                'total_pages': total_pages
                            }
                        )
                        self.documents.append(page_doc)
                        
                        # Split page into chunks
                        page_chunks = self.text_splitter.split_documents([page_doc])
                        self.chunks.extend(page_chunks)
                        
                        # Show progress for large documents
                        if total_pages > 20:
                            if (page_num + 1) % 10 == 0 or page_num + 1 == total_pages:
                                logger.info(f"   Processed {page_num + 1}/{total_pages} pages ({len(self.chunks)} chunks)")
                        else:
                            logger.info(f"   Page {page_num + 1}: {len(text)} characters, {len(page_chunks)} chunks")
            
            doc.close()
            logger.info(f"PDF processing completed: {len(self.documents)} page documents, {len(self.chunks)} total chunks")
            
        except Exception as e:
            logger.error(f"Failed to extract PDF content: {e}")
            return
    
    def _extract_docx_content(self, path: str) -> str:
        """Extract text from DOCX with enhanced formatting"""
        if DocxDocument is None:
            logger.error("python-docx not available. Install with: pip install python-docx")
            return ""
            
        try:
            doc = DocxDocument(path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Failed to extract DOCX content: {e}")
            return ""
    
    def _create_embeddings(self):
        """Create embeddings for all chunks using multiple models with memory optimization"""
        if not self.embedding_models:
            logger.warning("No embedding models available. Cannot create embeddings.")
            return
            
        if faiss is None:
            logger.warning("FAISS not available. Install with: pip install faiss-cpu")
            return
            
        chunk_texts = [chunk.page_content for chunk in self.chunks]
        total_chunks = len(chunk_texts)
        
        logger.info(f"Creating embeddings for {total_chunks} chunks...")
        
        for model_name, model in self.embedding_models.items():
            try:
                logger.info(f"Processing model: {model_name}")
                
                # Process in batches for large documents
                batch_size = 50  # Process 50 chunks at a time
                all_embeddings = []
                
                for i in range(0, total_chunks, batch_size):
                    batch = chunk_texts[i:i + batch_size]
                    batch_embeddings = model.encode(batch, show_progress_bar=False)
                    all_embeddings.append(batch_embeddings)
                    
                    # Show progress
                    if total_chunks > 100:
                        progress = min((i + batch_size) / total_chunks * 100, 100)
                        logger.info(f"   Progress: {progress:.1f}% ({i + len(batch)}/{total_chunks} chunks)")
                
                # Combine all embeddings
                embeddings = np.vstack(all_embeddings)
                
                # Create FAISS index
                dimension = embeddings.shape[1]
                index = faiss.IndexFlatIP(dimension)  # Inner product for cosine similarity
                
                # Normalize embeddings for cosine similarity
                faiss.normalize_L2(embeddings)
                index.add(embeddings.astype(np.float32))
                
                self.vector_stores[model_name] = {
                    'index': index,
                    'embeddings': embeddings
                }
                
                logger.info(f"Created FAISS index for {model_name} with {total_chunks} chunks")
                
            except Exception as e:
                logger.error(f"Failed to create embeddings for {model_name}: {e}")
                # Continue with other models if one fails
                continue
    
    def _build_keyword_index(self):
        """Build keyword index for hybrid search"""
        logger.info("Building keyword index...")
        
        for i, chunk in enumerate(self.chunks):
            content_lower = chunk.page_content.lower()
            keyword_matches = []
            
            for keyword in self.coverage_keywords:
                if keyword in content_lower:
                    keyword_matches.append(keyword)
            
            if keyword_matches:
                self.keyword_index[i] = keyword_matches
        
        logger.info(f"Built keyword index with {len(self.keyword_index)} entries")
    
    def retrieve_relevant_chunks(self, query: str, top_k: int = 20) -> List[RetrievedClause]:
        """Retrieve relevant chunks using advanced hybrid search with better context understanding"""
        if not self.embedding_models:
            logger.warning("No embedding models available. Returning empty results.")
            return []
            
        if not self.vector_stores:
            logger.warning("No vector stores available. Documents may not be loaded. Returning empty results.")
            return []
            
        if faiss is None:
            logger.warning("FAISS not available. Cannot perform similarity search.")
            return []
            
        all_results = []
        
        # Enhanced query processing with domain-specific terms
        query_lower = query.lower()
        query_terms = query_lower.split()
        
        # Domain-specific term categories for better matching
        insurance_terms = ['policy', 'insurance', 'coverage', 'claim', 'premium', 'benefit', 'exclusion']
        medical_terms = ['surgery', 'treatment', 'procedure', 'medical', 'hospital', 'doctor', 'patient']
        time_terms = ['waiting', 'period', 'grace', 'month', 'year', 'days']
        amount_terms = ['amount', 'limit', 'maximum', 'coverage', 'sum', 'insured']
        specific_terms = ['organ', 'donor', 'maternity', 'cataract', 'ncd', 'ayush', 'room', 'rent']
        
        # Identify query type for better retrieval
        query_type = 'general'
        if any(term in query_lower for term in ['organ', 'donor']):
            query_type = 'organ_donor'
        elif any(term in query_lower for term in ['grace', 'period']):
            query_type = 'grace_period'
        elif any(term in query_lower for term in ['waiting', 'period']):
            query_type = 'waiting_period'
        elif any(term in query_lower for term in ['maternity', 'pregnancy']):
            query_type = 'maternity'
        elif any(term in query_lower for term in ['cataract']):
            query_type = 'cataract'
        elif any(term in query_lower for term in ['ncd', 'no claim']):
            query_type = 'ncd'
        elif any(term in query_lower for term in ['health check', 'preventive']):
            query_type = 'health_check'
        elif any(term in query_lower for term in ['hospital', 'definition']):
            query_type = 'hospital_definition'
        elif any(term in query_lower for term in ['ayush']):
            query_type = 'ayush'
        elif any(term in query_lower for term in ['room rent', 'sub-limit']):
            query_type = 'room_rent'
        
        # Hybrid search: Combine semantic and keyword search
        for model_name, model in self.embedding_models.items():
            if model_name not in self.vector_stores:
                continue
            
            # Encode query
            query_embedding = model.encode([query])
            faiss.normalize_L2(query_embedding)
            
            # Search with more results for better context
            vector_store = self.vector_stores[model_name]
            scores, indices = vector_store['index'].search(
                query_embedding.astype(np.float32), top_k * 5  # Get more results for filtering
            )
            
            # Collect results with enhanced scoring
            for score, idx in zip(scores[0], indices[0]):
                if idx < len(self.chunks):
                    chunk = self.chunks[idx]
                    content_lower = chunk.page_content.lower()
                    
                    # Calculate semantic similarity
                    semantic_similarity = float(score)
                    
                    # Calculate comprehensive keyword boost
                    keyword_boost = 0.0
                    keyword_matches = []
                    
                    # Boost for exact query term matches
                    for term in query_terms:
                        if len(term) > 2 and term in content_lower:
                            keyword_boost += 0.3
                            keyword_matches.append(term)
                    
                    # Boost for domain-specific terms
                    for term in insurance_terms:
                        if term in content_lower:
                            keyword_boost += 0.1
                            keyword_matches.append(term)
                    
                    for term in medical_terms:
                        if term in content_lower:
                            keyword_boost += 0.15
                            keyword_matches.append(term)
                    
                    for term in time_terms:
                        if term in content_lower:
                            keyword_boost += 0.1
                            keyword_matches.append(term)
                    
                    for term in amount_terms:
                        if term in content_lower:
                            keyword_boost += 0.1
                            keyword_matches.append(term)
                    
                    # High boost for specific terms related to query type
                    if query_type != 'general':
                        specific_boost_terms = {
                            'organ_donor': ['organ', 'donor', 'donation'],
                            'grace_period': ['grace', 'period', 'premium', 'payment'],
                            'waiting_period': ['waiting', 'period', 'pre-existing'],
                            'maternity': ['maternity', 'pregnancy', 'childbirth'],
                            'cataract': ['cataract'],
                            'ncd': ['ncd', 'no claim discount'],
                            'health_check': ['health check', 'preventive'],
                            'hospital_definition': ['hospital', 'definition', 'bed'],
                            'ayush': ['ayush', 'ayurveda', 'yoga'],
                            'room_rent': ['room rent', 'sub-limit', 'icu']
                        }
                        
                        if query_type in specific_boost_terms:
                            for term in specific_boost_terms[query_type]:
                                if term in content_lower:
                                    keyword_boost += 0.4  # High boost for specific terms
                                    keyword_matches.append(term)
                    
                    # Boost for coverage-related terms
                    coverage_terms = ['covered', 'eligible', 'approved', 'excluded', 'not covered', 'yes', 'no']
                    for term in coverage_terms:
                        if term in content_lower:
                            keyword_boost += 0.05
                            keyword_matches.append(term)
                    
                    # Calculate final score with weighted combination
                    final_score = (semantic_similarity * 0.6) + (keyword_boost * 0.4)
                    
                    clause = RetrievedClause(
                        content=chunk.page_content,
                        source_document=chunk.metadata.get('source', ''),
                        relevance_score=final_score,
                        clause_type=model_name,
                        semantic_similarity=semantic_similarity,
                        keyword_matches=list(set(keyword_matches))  # Remove duplicates
                    )
                    all_results.append(clause)
        
        # Remove duplicates and sort by relevance
        unique_results = {}
        for result in all_results:
            key = result.content[:400]  # Use first 400 chars as key for better deduplication
            if key not in unique_results or result.relevance_score > unique_results[key].relevance_score:
                unique_results[key] = result
        
        # Sort by relevance score
        final_results = sorted(unique_results.values(), 
                             key=lambda x: x.relevance_score, reverse=True)
        
        return final_results[:top_k]

class AdvancedLLMDecisionEngine:
    """Advanced decision engine using powerful LLM for intelligent question answering"""
    
    def __init__(self, model_name: str = "microsoft/DialoGPT-large"):
        self.model_name = model_name
        self.llm_available = False
        self.tokenizer = None
        self.model = None
        self.qa_pipeline = None
        
        # Initialize multiple LLM components for different tasks
        if AutoTokenizer is not None and AutoModel is not None and pipeline is not None:
            try:
                logger.info(f"Loading powerful LLM model: {model_name}")
                
                # Load tokenizer and model for text generation
                self.tokenizer = AutoTokenizer.from_pretrained(model_name)
                self.model = AutoModel.from_pretrained(model_name)
                
                # Load a more sophisticated QA model
                logger.info("Loading QA pipeline...")
                self.qa_pipeline = pipeline(
                    "question-answering",
                    model="deepset/roberta-large-squad2",  # More powerful QA model
                    tokenizer="deepset/roberta-large-squad2"
                )
                
                self.llm_available = True
                logger.info("Powerful LLM models loaded successfully")
                
                # Test the LLM with a simple question
                try:
                    test_answer = self.answer_question_with_llm("What is a grace period?", "A grace period is a time period after the due date during which payment can be made.")
                    logger.info(f"LLM test successful: {test_answer[:50]}...")
                except Exception as e:
                    logger.warning(f"LLM test failed: {e}")
                    self.llm_available = False
            except Exception as e:
                logger.warning(f"Failed to load LLM models: {e}")
                self.llm_available = False
        else:
            logger.warning("Required LLM libraries not available")
            self.llm_available = False
        
        # Enhanced policy rules and decision logic (fallback)
        self.policy_rules = {
            'age_limits': {
                'knee_surgery': {'min': 18, 'max': 65},
                'cardiac_surgery': {'min': 18, 'max': 70},
                'dental_surgery': {'min': 18, 'max': 75},
                'general_surgery': {'min': 18, 'max': 80}
            },
            'waiting_periods': {
                'knee_surgery': 6,
                'cardiac_surgery': 12,
                'dental_surgery': 3,
                'general_surgery': 6
            },
            'coverage_amounts': {
                'knee_surgery': 50000,
                'cardiac_surgery': 100000,
                'dental_surgery': 3000,
                'general_surgery': 25000
            }
        }
        
    def make_decision(self, query: QueryStructured, 
                     relevant_clauses: List[RetrievedClause]) -> DecisionResponse:
        """Make decision using LLM-based question answering"""
        
        # Create enhanced context from relevant clauses
        context = self._create_enhanced_context(relevant_clauses)
        
        # Always try LLM first, regardless of availability
        try:
            if self.llm_available:
                logger.info("Using LLM-based question answering")
                decision = self._answer_with_llm(query, context, relevant_clauses)
            else:
                logger.warning("LLM not available, using fallback method")
                # Fallback to rule-based decision
                reasoning_chain = self._generate_detailed_reasoning_chain(query, context)
                decision = self._evaluate_decision_enhanced(query, context, reasoning_chain)
        except Exception as e:
            logger.error(f"Error in LLM decision making: {e}")
            # Fallback to rule-based decision
            reasoning_chain = self._generate_detailed_reasoning_chain(query, context)
            decision = self._evaluate_decision_enhanced(query, context, reasoning_chain)
        
        # Add relevant clauses to decision
        decision.relevant_clauses = relevant_clauses[:5]  # Top 5 most relevant clauses
        
        return decision
    
    def answer_question_with_llm(self, question: str, context: str) -> str:
        """Answer a question using powerful LLM with context understanding"""
        if not self.llm_available:
            return "LLM not available for question answering."
        
        try:
            # Method 1: Use sophisticated QA pipeline
            if self.qa_pipeline:
                result = self.qa_pipeline(
                    question=question,
                    context=context,
                    max_answer_len=300,
                    handle_impossible_answer=True,
                    top_k=3  # Get multiple answers
                )
                
                # If we have a good answer from QA pipeline
                if isinstance(result, dict) and result.get('score', 0) > 0.4:
                    return result['answer']
                elif isinstance(result, list) and len(result) > 0:
                    # Return the best answer from multiple results
                    best_answer = result[0]
                    if best_answer.get('score', 0) > 0.4:
                        return best_answer['answer']
            
            # Method 2: Use text generation with context
            if self.tokenizer and self.model:
                return self._generate_contextual_answer(question, context)
            
            return "The answer to this question is not clearly stated in the provided policy document."
                
        except Exception as e:
            logger.error(f"Error in LLM question answering: {e}")
            return f"Error processing question: {str(e)}"
    
    def _generate_contextual_answer(self, question: str, context: str) -> str:
        """Generate contextual answer using text generation model"""
        try:
            # Create a prompt that includes context and question
            prompt = f"""Based on the following insurance policy document, please answer the question accurately and completely.

Policy Document:
{context[:2000]}  # Limit context length

Question: {question}

Answer:"""
            
            # Tokenize the prompt
            inputs = self.tokenizer.encode(prompt, return_tensors="pt", max_length=512, truncation=True)
            
            # Generate answer
            with torch.no_grad():
                outputs = self.model.generate(
                    inputs,
                    max_length=inputs.shape[1] + 200,  # Generate up to 200 more tokens
                    num_return_sequences=1,
                    temperature=0.7,
                    do_sample=True,
                    pad_token_id=self.tokenizer.eos_token_id
                )
            
            # Decode the generated text
            generated_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # Extract the answer part (after "Answer:")
            if "Answer:" in generated_text:
                answer = generated_text.split("Answer:")[-1].strip()
                return answer if answer else "Unable to generate a specific answer."
            else:
                return generated_text.split("Question:")[-1].strip() if "Question:" in generated_text else "Unable to generate a specific answer."
                
        except Exception as e:
            logger.error(f"Error in contextual answer generation: {e}")
            return "Error generating contextual answer."
    
    def _create_enhanced_context(self, clauses: List[RetrievedClause]) -> str:
        """Create enhanced structured context from retrieved clauses with better organization"""
        if not clauses:
            return "No relevant information found in the policy document."
        
        # Sort clauses by relevance score
        sorted_clauses = sorted(clauses, key=lambda x: x.relevance_score, reverse=True)
        
        # Create a comprehensive context
        context_parts = []
        context_parts.append("POLICY DOCUMENT CONTEXT:")
        context_parts.append("=" * 50)
        
        # Add the most relevant clauses (up to 5)
        for i, clause in enumerate(sorted_clauses[:5]):
            context_parts.append(f"\nRELEVANT SECTION {i+1} (Relevance: {clause.relevance_score:.3f}):")
            context_parts.append("-" * 30)
            
            # Clean and format the content
            content = clause.content.strip()
            if len(content) > 500:
                # Truncate very long content but keep important parts
                content = content[:500] + "..."
            
            context_parts.append(content)
            
            # Add keyword information if available
            if clause.keyword_matches:
                context_parts.append(f"\nKey Terms: {', '.join(clause.keyword_matches)}")
            
            context_parts.append(f"Source: {clause.source_document}")
        
        # Add summary of all relevant information
        context_parts.append("\n" + "=" * 50)
        context_parts.append("SUMMARY:")
        
        # Extract key information from all clauses
        all_content = " ".join([c.content for c in sorted_clauses[:3]])
        
        # Look for specific patterns in the content
        patterns = {
            'grace_period': r'grace\s+period.*?(\d+)\s*(?:days?|months?)',
            'waiting_period': r'waiting\s+period.*?(\d+)\s*(?:months?|years?)',
            'coverage_amount': r'(?:coverage|amount|up\s+to).*?(\d+(?:,\d+)*(?:\.\d+)?)',
            'organ_donor': r'organ\s+(?:donor|donation)',
            'maternity': r'maternity|pregnancy|childbirth',
            'cataract': r'cataract',
            'ncd': r'ncd|no\s+claim\s+discount',
            'health_check': r'health\s+check|preventive',
            'hospital': r'hospital.*?(?:definition|defined)',
            'ayush': r'ayush|ayurveda|yoga|naturopathy',
            'room_rent': r'room\s+rent|sub.?limit'
        }
        
        found_info = []
        for pattern_name, pattern in patterns.items():
            matches = re.findall(pattern, all_content, re.IGNORECASE)
            if matches:
                found_info.append(f"{pattern_name.replace('_', ' ').title()}: {', '.join(matches)}")
        
        if found_info:
            context_parts.append("Key Information Found:")
            for info in found_info:
                context_parts.append(f"• {info}")
        
        return "\n".join(context_parts)
    
    def _generate_detailed_reasoning_chain(self, query: QueryStructured, context: str) -> List[str]:
        """Generate detailed step-by-step reasoning chain"""
        reasoning_steps = []
        
        # Step 1: Query analysis
        reasoning_steps.append(f"Query Analysis: {query.age}Y {query.gender}, {query.procedure} in {query.location}, {query.policy_duration} policy")
        
        # Step 2: Policy eligibility check
        if query.policy_age_months:
            if query.policy_age_months < 6:
                reasoning_steps.append(f"Policy Age Check: Policy is {query.policy_age_months} months old (may have waiting period restrictions)")
            elif query.policy_age_months < 12:
                reasoning_steps.append(f"Policy Age Check: Policy is {query.policy_age_months} months old (standard coverage likely)")
            else:
                reasoning_steps.append(f"Policy Age Check: Policy is {query.policy_age_months} months old (full coverage available)")
        
        # Step 3: Procedure coverage check
        if query.procedure:
            reasoning_steps.append(f"Procedure Check: Evaluating coverage for {query.procedure}")
        
        # Step 4: Location coverage check
        if query.location:
            reasoning_steps.append(f"Location Check: Treatment location is {query.location}")
        
        # Step 5: Age-based considerations
        if query.age:
            if query.age > 65:
                reasoning_steps.append(f"Age Check: Senior citizen ({query.age} years) - may have age-specific restrictions")
            elif query.age < 18:
                reasoning_steps.append(f"Age Check: Minor ({query.age} years) - may have guardian requirements")
            else:
                reasoning_steps.append(f"Age Check: Adult ({query.age} years) - standard coverage applies")
        
        # Step 6: Context evaluation
        reasoning_steps.append("Context Evaluation: Analyzing relevant policy clauses for coverage determination")
        
        return reasoning_steps
    
    def _evaluate_decision_enhanced(self, query: QueryStructured, context: str, 
                                   reasoning_chain: List[str]) -> DecisionResponse:
        """Evaluate final decision based on enhanced document content analysis"""
        
        decision = "requires_review"  # Default
        amount = None
        confidence = 0.5
        justification = "Manual review required"
        coverage_details = {}
        
        context_lower = context.lower()
        query_procedure = query.procedure.lower() if query.procedure else ""
        
        # Enhanced coverage analysis with specific policy rules
        coverage_indicators = {
            'covered': 0,
            'eligible': 0,
            'approved': 0,
            'excluded': 0,
            'not covered': 0,
            'waiting period': 0,
            'pre-existing': 0,
            'knee': 0,
            'surgery': 0,
            'treatment': 0,
            'maximum': 0,
            'limit': 0,
            'coverage': 0
        }
        
        # Count coverage indicators in context
        for indicator in coverage_indicators:
            coverage_indicators[indicator] = context_lower.count(indicator)
        
        # Enhanced procedure analysis
        procedure_mentioned = query_procedure in context_lower if query_procedure else False
        surgery_mentioned = 'surgery' in context_lower
        knee_mentioned = 'knee' in context_lower
        
        # Check for specific procedure mentions with context
        procedure_coverage = self._analyze_procedure_coverage(query_procedure, context_lower)
        
        # Enhanced decision logic with policy rules
        if procedure_coverage['explicitly_covered']:
            decision = "approved"
            confidence = 0.9
            justification = f"{query.procedure} is explicitly covered in the policy"
            coverage_details = procedure_coverage
        elif procedure_coverage['explicitly_excluded']:
            decision = "rejected"
            confidence = 0.85
            justification = f"{query.procedure} is explicitly excluded from coverage"
            coverage_details = procedure_coverage
        elif procedure_coverage['conditionally_covered']:
            decision = "approved"
            confidence = 0.75
            justification = f"{query.procedure} is covered with conditions"
            coverage_details = procedure_coverage
        elif surgery_mentioned and coverage_indicators['covered'] > 0:
            decision = "approved"
            confidence = 0.8
            justification = "Surgery is covered under the policy"
        elif knee_mentioned and 'knee' in query_procedure and coverage_indicators['covered'] > 0:
            decision = "approved"
            confidence = 0.85
            justification = "Knee surgery is covered under the policy"
        else:
            # Apply policy rules for decision making
            decision_result = self._apply_policy_rules(query, context_lower)
            decision = decision_result['decision']
            confidence = decision_result['confidence']
            justification = decision_result['justification']
            coverage_details = decision_result['coverage_details']
        
        # Extract amounts with enhanced patterns
        amount = self._extract_coverage_amount(context_lower)
        
        # Policy age considerations
        if query.policy_age_months and query.policy_age_months < 6:
            if coverage_indicators['waiting period'] > 0:
                confidence *= 0.8
                justification += " (Waiting period may apply)"
        
        # Age-based adjustments
        if query.age and query.age > 65:
            confidence *= 0.9
            justification += " (Senior citizen considerations)"
        
        return DecisionResponse(
            decision=decision,
            amount=amount,
            confidence=confidence,
            justification=justification,
            relevant_clauses=None,
            reasoning_chain=reasoning_chain,
            coverage_details=coverage_details
        )
    
    def _apply_policy_rules(self, query: QueryStructured, context: str) -> Dict[str, Any]:
        """Apply specific policy rules for decision making"""
        result = {
            'decision': 'requires_review',
            'confidence': 0.6,
            'justification': 'Coverage status unclear - manual review required',
            'coverage_details': {}
        }
        
        if not query.procedure:
            return result
        
        procedure_lower = query.procedure.lower()
        
        # Determine procedure type
        procedure_type = None
        if 'knee' in procedure_lower:
            procedure_type = 'knee_surgery'
        elif 'cardiac' in procedure_lower or 'heart' in procedure_lower:
            procedure_type = 'cardiac_surgery'
        elif 'dental' in procedure_lower:
            procedure_type = 'dental_surgery'
        else:
            procedure_type = 'general_surgery'
        
        # Check age limits
        age_limits = {'min': 18, 'max': 80}
        if 'knee' in procedure_lower:
            age_limits = {'min': 18, 'max': 65}
        elif 'cardiac' in procedure_lower or 'heart' in procedure_lower:
            age_limits = {'min': 18, 'max': 70}
        
        if query.age:
            if query.age < age_limits['min'] or query.age > age_limits['max']:
                result['decision'] = 'rejected'
                result['confidence'] = 0.9
                result['justification'] = f"Age {query.age} is outside coverage limits ({age_limits['min']}-{age_limits['max']} years) for {procedure_type}"
                return result
        
        # Check waiting period
        waiting_period = 6  # Default
        if 'knee' in procedure_lower:
            waiting_period = 6
        elif 'cardiac' in procedure_lower or 'heart' in procedure_lower:
            waiting_period = 12
        elif 'dental' in procedure_lower:
            waiting_period = 3
        
        if query.policy_age_months and query.policy_age_months < waiting_period:
            result['decision'] = 'rejected'
            result['confidence'] = 0.85
            result['justification'] = f"Policy is {query.policy_age_months} months old, but {waiting_period}-month waiting period applies for {procedure_type}"
            return result
        
        # Check if procedure is mentioned in context
        if procedure_type in context or query.procedure.lower() in context:
            result['decision'] = 'approved'
            result['confidence'] = 0.8
            result['justification'] = f"{query.procedure} is covered under the policy"
            
            # Add coverage amount
            coverage_amount = 25000  # Default
            if 'knee' in procedure_lower:
                coverage_amount = 50000
            elif 'cardiac' in procedure_lower or 'heart' in procedure_lower:
                coverage_amount = 100000
            elif 'dental' in procedure_lower:
                coverage_amount = 3000
            
            result['coverage_details'] = {
                'coverage_amount': coverage_amount,
                'procedure_type': procedure_type,
                'waiting_period': waiting_period,
                'age_limits': age_limits
            }
        else:
            # Check for general surgery coverage
            if 'surgery' in context and ('covered' in context or 'eligible' in context):
                result['decision'] = 'approved'
                result['confidence'] = 0.75
                result['justification'] = f"General surgery coverage applies to {query.procedure}"
            else:
                result['decision'] = 'requires_review'
                result['confidence'] = 0.6
                result['justification'] = f"Specific coverage for {query.procedure} not found in policy documents"
        
        return result
    
    def _analyze_procedure_coverage(self, procedure: str, context: str) -> Dict[str, Any]:
        """Analyze specific procedure coverage from context"""
        analysis = {
            'explicitly_covered': False,
            'explicitly_excluded': False,
            'conditionally_covered': False,
            'coverage_conditions': [],
            'waiting_period': None,
            'maximum_amount': None
        }
        
        if not procedure:
            return analysis
        
        # Check for explicit coverage
        coverage_patterns = [
            f"{procedure}.*covered",
            f"{procedure}.*eligible",
            f"{procedure}.*approved",
            f"covered.*{procedure}",
            f"eligible.*{procedure}"
        ]
        
        for pattern in coverage_patterns:
            if re.search(pattern, context, re.IGNORECASE):
                analysis['explicitly_covered'] = True
                break
        
        # Check for explicit exclusion
        exclusion_patterns = [
            f"{procedure}.*excluded",
            f"{procedure}.*not covered",
            f"excluded.*{procedure}",
            f"not covered.*{procedure}"
        ]
        
        for pattern in exclusion_patterns:
            if re.search(pattern, context, re.IGNORECASE):
                analysis['explicitly_excluded'] = True
                break
        
        # Check for conditional coverage
        conditional_patterns = [
            f"{procedure}.*condition",
            f"{procedure}.*subject to",
            f"{procedure}.*provided",
            f"condition.*{procedure}"
        ]
        
        for pattern in conditional_patterns:
            if re.search(pattern, context, re.IGNORECASE):
                analysis['conditionally_covered'] = True
                break
        
        # Extract waiting period
        waiting_match = re.search(r'(\d+)\s*(?:month|mon|m).*waiting', context, re.IGNORECASE)
        if waiting_match:
            analysis['waiting_period'] = int(waiting_match.group(1))
        
        # Extract maximum amount
        amount_match = re.search(r'(\d+(?:,\d+)*(?:\.\d+)?).*maximum', context, re.IGNORECASE)
        if amount_match:
            try:
                analysis['maximum_amount'] = float(amount_match.group(1).replace(',', ''))
            except ValueError:
                pass
        
        return analysis
    
    def _answer_with_llm(self, query: QueryStructured, context: str, relevant_clauses: List[RetrievedClause]) -> DecisionResponse:
        """Answer questions using LLM based on retrieved context"""
        
        # Convert structured query back to natural language question
        question = query.raw_query if query.raw_query else "What is the coverage for this procedure?"
        
        logger.info(f"LLM Question: {question}")
        logger.info(f"LLM Context length: {len(context)} characters")
        
        # Use LLM to answer the question
        llm_answer = self.answer_question_with_llm(question, context)
        
        logger.info(f"LLM Answer: {llm_answer}")
        
        # Determine decision based on LLM answer
        answer_lower = llm_answer.lower()
        
        if any(word in answer_lower for word in ['yes', 'covered', 'eligible', 'approved']):
            decision = "approved"
            confidence = 0.8
        elif any(word in answer_lower for word in ['no', 'not covered', 'excluded', 'not eligible']):
            decision = "rejected"
            confidence = 0.8
        else:
            decision = "requires_review"
            confidence = 0.6
        
        # Extract amount if mentioned
        amount = self._extract_coverage_amount(context)
        
        return DecisionResponse(
            decision=decision,
            amount=amount,
            confidence=confidence,
            justification=llm_answer,
            relevant_clauses=relevant_clauses,
            reasoning_chain=[f"LLM analysis: {llm_answer}"],
            coverage_details={"llm_answer": llm_answer}
        )
    
    def _extract_coverage_amount(self, context: str) -> Optional[float]:
        """Extract coverage amount from context with enhanced patterns"""
        amount_patterns = [
            r'coverage.*?(\d+(?:,\d+)*(?:\.\d+)?)',
            r'amount.*?(\d+(?:,\d+)*(?:\.\d+)?)',
            r'up to.*?(\d+(?:,\d+)*(?:\.\d+)?)',
            r'maximum.*?(\d+(?:,\d+)*(?:\.\d+)?)',
            r'limit.*?(\d+(?:,\d+)*(?:\.\d+)?)',
            r'(\d+(?:,\d+)*(?:\.\d+)?).*coverage',
            r'(\d+(?:,\d+)*(?:\.\d+)?).*maximum'
        ]
        
        for pattern in amount_patterns:
            amount_match = re.search(pattern, context, re.IGNORECASE)
            if amount_match:
                try:
                    amount = float(amount_match.group(1).replace(',', ''))
                    return amount
                except ValueError:
                    continue
        
        return None

class EnhancedRAGSystem:
    """Enhanced RAG system with improved accuracy and information retrieval"""
    
    def __init__(self):
        self.query_parser = EnhancedQueryParser()
        self.retriever = MultiModelEmbeddingRetriever()
        self.decision_engine = AdvancedLLMDecisionEngine()
        
        logger.info("Enhanced RAG System initialized")
    
    def load_documents(self, document_paths: List[str]):
        """Load documents into the system"""
        self.retriever.load_documents(document_paths)
    
    def process_query(self, query: str, debug: bool = False) -> Dict[str, Any]:
        """Process a natural language query and return structured decision with enhanced accuracy"""
        
        # Step 1: Parse query with enhanced extraction
        structured_query = self.query_parser.parse_query(query)
        logger.info(f"Parsed query: {structured_query}")
        
        # Step 2: Retrieve relevant information with hybrid search
        relevant_clauses = self.retriever.retrieve_relevant_chunks(query, top_k=15)
        logger.info(f"Retrieved {len(relevant_clauses)} relevant clauses")
        
        # Debug: Show retrieved content with enhanced details
        if debug and relevant_clauses:
            print(f"\nDEBUG: Retrieved Content:")
            for i, clause in enumerate(relevant_clauses[:5], 1):
                print(f"   Clause {i} (Score: {clause.relevance_score:.3f}, Semantic: {clause.semantic_similarity:.3f}):")
                content_preview = clause.content[:200].replace('\n', ' ').strip()
                print(f"      {content_preview}...")
                if clause.keyword_matches:
                    print(f"      Keywords: {', '.join(clause.keyword_matches)}")
        elif debug:
            print(f"\nDEBUG: No relevant clauses found!")
            print(f"   Total chunks available: {len(self.retriever.chunks)}")
            print(f"   Query: '{query}'")
            
            # Show some sample chunks to see what's in the document
            if self.retriever.chunks:
                print(f"\nSample document chunks:")
                for i, chunk in enumerate(self.retriever.chunks[:3], 1):
                    preview = chunk.page_content[:100].replace('\n', ' ').strip()
                    print(f"   Chunk {i}: {preview}...")
        
        # Step 3: Make decision with enhanced analysis
        decision = self.decision_engine.make_decision(structured_query, relevant_clauses)
        logger.info(f"Decision: {decision.decision} (confidence: {decision.confidence:.2f})")
        
        # Step 4: Format response with enhanced structure for insurance processing
        response = {
            "decision": decision.decision,
            "amount": decision.amount,
            "confidence": decision.confidence,
            "justification": decision.justification,
            "reasoning_chain": decision.reasoning_chain,
            "coverage_details": decision.coverage_details,
            "query_analysis": {
                "age": structured_query.age,
                "gender": structured_query.gender,
                "procedure": structured_query.procedure,
                "location": structured_query.location,
                "policy_duration": structured_query.policy_duration,
                "policy_age_months": structured_query.policy_age_months,
                "extracted_entities": structured_query.extracted_entities
            },
            "relevant_clauses": [
                {
                    "content": clause.content[:400] + "..." if len(clause.content) > 400 else clause.content,
                    "source": clause.source_document,
                    "relevance_score": clause.relevance_score,
                    "semantic_similarity": clause.semantic_similarity,
                    "keyword_matches": clause.keyword_matches,
                    "clause_type": clause.clause_type,
                    "clause_id": f"clause_{i+1}"
                }
                for i, clause in enumerate(relevant_clauses[:8])  # Top 8 clauses
            ],
            "policy_mapping": {
                "age_compliance": self._check_age_compliance(structured_query),
                "policy_age_compliance": self._check_policy_age_compliance(structured_query),
                "location_coverage": self._check_location_coverage(structured_query),
                "procedure_coverage": self._check_procedure_coverage(structured_query, relevant_clauses)
            },
            "summary": {
                "total_clauses_retrieved": len(relevant_clauses),
                "top_relevance_score": max([c.relevance_score for c in relevant_clauses]) if relevant_clauses else 0,
                "average_relevance_score": sum([c.relevance_score for c in relevant_clauses]) / len(relevant_clauses) if relevant_clauses else 0,
                "high_relevance_clauses": len([c for c in relevant_clauses if c.relevance_score > 0.7]),
                "keyword_matches_found": len(set([kw for c in relevant_clauses if c.keyword_matches for kw in c.keyword_matches])),
                "decision_basis": self._get_decision_basis(decision, relevant_clauses)
            }
        }
        
        return response
    
    def _check_age_compliance(self, query: QueryStructured) -> Dict[str, Any]:
        """Check if age meets policy requirements"""
        if not query.age:
            return {"status": "unknown", "reason": "Age not specified"}
        
        # Default age limits
        min_age, max_age = 18, 80
        
        if query.procedure and 'knee' in query.procedure.lower():
            max_age = 65
        elif query.procedure and ('cardiac' in query.procedure.lower() or 'heart' in query.procedure.lower()):
            max_age = 70
        
        if query.age < min_age:
            return {"status": "non_compliant", "reason": f"Age {query.age} below minimum {min_age}"}
        elif query.age > max_age:
            return {"status": "non_compliant", "reason": f"Age {query.age} above maximum {max_age}"}
        else:
            return {"status": "compliant", "reason": f"Age {query.age} within limits {min_age}-{max_age}"}
    
    def _check_policy_age_compliance(self, query: QueryStructured) -> Dict[str, Any]:
        """Check if policy age meets requirements"""
        if not query.policy_age_months:
            return {"status": "unknown", "reason": "Policy age not specified"}
        
        # Default waiting period
        waiting_period = 6
        
        if query.procedure and 'knee' in query.procedure.lower():
            waiting_period = 6
        elif query.procedure and ('cardiac' in query.procedure.lower() or 'heart' in query.procedure.lower()):
            waiting_period = 12
        elif query.procedure and 'dental' in query.procedure.lower():
            waiting_period = 3
        
        if query.policy_age_months < waiting_period:
            return {"status": "non_compliant", "reason": f"Policy age {query.policy_age_months} months, waiting period {waiting_period} months"}
        else:
            return {"status": "compliant", "reason": f"Policy age {query.policy_age_months} months meets waiting period {waiting_period} months"}
    
    def _check_location_coverage(self, query: QueryStructured) -> Dict[str, Any]:
        """Check if location is covered"""
        if not query.location:
            return {"status": "unknown", "reason": "Location not specified"}
        
        covered_cities = ['pune', 'mumbai', 'delhi', 'bangalore', 'chennai', 'hyderabad', 
                         'kolkata', 'ahmedabad', 'nagpur', 'indore', 'bhopal', 'lucknow']
        
        if query.location.lower() in covered_cities:
            return {"status": "covered", "reason": f"Location {query.location} is in network"}
        else:
            return {"status": "unknown", "reason": f"Location {query.location} coverage unclear"}
    
    def _check_procedure_coverage(self, query: QueryStructured, clauses: List[RetrievedClause]) -> Dict[str, Any]:
        """Check if procedure is covered based on retrieved clauses"""
        if not query.procedure:
            return {"status": "unknown", "reason": "Procedure not specified"}
        
        # Check if procedure is mentioned in any clause
        procedure_lower = query.procedure.lower()
        for clause in clauses:
            if procedure_lower in clause.content.lower():
                return {"status": "covered", "reason": f"Procedure {query.procedure} found in policy documents"}
        
        return {"status": "unknown", "reason": f"Procedure {query.procedure} not found in policy documents"}
    
    def _get_decision_basis(self, decision: DecisionResponse, clauses: List[RetrievedClause]) -> str:
        """Get the basis for the decision"""
        if decision.decision == "approved":
            return "Policy coverage confirmed based on retrieved clauses"
        elif decision.decision == "rejected":
            return "Policy exclusion or non-compliance identified"
        else:
            return "Insufficient information for automatic decision"
    
    def batch_process(self, queries: List[str]) -> List[Dict[str, Any]]:
        """Process multiple queries in batch with enhanced accuracy"""
        results = []
        for query in queries:
            try:
                result = self.process_query(query)
                results.append(result)
            except Exception as e:
                logger.error(f"Error processing query '{query}': {e}")
                results.append({
                    "error": str(e),
                    "query": query
                })
        return results

# Example usage and testing
def create_sample_documents():
    """Create sample insurance policy documents for testing"""
    sample_docs = [
        {
            "content": """
            NATIONAL PARIVAR MEDICLAIM PLUS POLICY
            
            SECTION 1: PREMIUM PAYMENT AND GRACE PERIOD
            
            Grace Period for Premium Payment:
            A grace period of thirty days is provided for premium payment after the due date to renew or continue the policy without losing continuity benefits. During this grace period, the policy remains active and coverage continues.
            
            SECTION 2: PRE-EXISTING DISEASES (PED)
            
            Waiting Period for Pre-existing Diseases:
            There is a waiting period of thirty-six (36) months of continuous coverage from the first policy inception for pre-existing diseases and their direct complications to be covered. This waiting period applies to all conditions that existed before the policy start date.
            
            SECTION 3: MATERNITY COVERAGE
            
            Maternity Expenses Coverage:
            Yes, the policy covers maternity expenses, including childbirth and lawful medical termination of pregnancy. To be eligible, the female insured person must have been continuously covered for at least 24 months. The benefit is limited to two deliveries or terminations during the policy period.
            
            SECTION 4: CATARACT SURGERY
            
            Cataract Surgery Waiting Period:
            The policy has a specific waiting period of two (2) years for cataract surgery. This waiting period must be completed before any cataract-related procedures are covered.
            
            SECTION 5: ORGAN DONOR EXPENSES
            
            Organ Donor Medical Expenses:
            Yes, the policy indemnifies the medical expenses for the organ donor's hospitalization for the purpose of harvesting the organ, provided the organ is for an insured person and the donation complies with the Transplantation of Human Organs Act, 1994.
            
            SECTION 6: NO CLAIM DISCOUNT (NCD)
            
            No Claim Discount Benefits:
            A No Claim Discount of 5% on the base premium is offered on renewal for a one-year policy term if no claims were made in the preceding year. The maximum aggregate NCD is capped at 5% of the total base premium.
            
            SECTION 7: PREVENTIVE HEALTH CHECK-UPS
            
            Preventive Health Check-up Benefits:
            Yes, the policy reimburses expenses for health check-ups at the end of every block of two continuous policy years, provided the policy has been renewed without a break. The amount is subject to the limits specified in the Table of Benefits.
            
            SECTION 8: HOSPITAL DEFINITION
            
            Definition of Hospital:
            A hospital is defined as an institution with at least 10 inpatient beds (in towns with a population below ten lakhs) or 15 beds (in all other places), with qualified nursing staff and medical practitioners available 24/7, a fully equipped operation theatre, and which maintains daily records of patients.
            
            SECTION 9: AYUSH TREATMENTS
            
            AYUSH Coverage:
            The policy covers medical expenses for inpatient treatment under Ayurveda, Yoga, Naturopathy, Unani, Siddha, and Homeopathy systems up to the Sum Insured limit, provided the treatment is taken in an AYUSH Hospital.
            
            SECTION 10: ROOM RENT AND ICU CHARGES
            
            Room Rent and ICU Sub-limits for Plan A:
            Yes, for Plan A, the daily room rent is capped at 1% of the Sum Insured, and ICU charges are capped at 2% of the Sum Insured. These limits do not apply if the treatment is for a listed procedure in a Preferred Provider Network (PPN).
            
            SECTION 11: GENERAL COVERAGE
            
            Knee Surgery Coverage:
            - Knee surgery is covered for patients aged 18-65 years
            - Waiting period: 6 months for major knee surgeries
            - Maximum coverage amount: $50,000 for knee procedures
            - Pre-authorization required for surgeries over $10,000
            - Covers: knee replacement, arthroscopy, ligament repair
            - Network hospitals in Pune, Mumbai, Delhi, Bangalore
            
            Age Restrictions:
            - Minimum age: 18 years
            - Maximum age: 65 years for knee procedures
            - Senior citizens (65+) require additional approval
            
            Waiting Periods:
            - 6 months for major orthopedic surgeries
            - 3 months for minor procedures
            - Pre-existing conditions: 12-month waiting period
            
            Network Coverage:
            - All major hospitals in Pune, Mumbai, Delhi, Bangalore
            - Cashless treatment available at network hospitals
            - Emergency procedures covered immediately
            """,
            "source": "national_parivar_mediclaim_plus_policy.txt"
        },
        {
            "content": """
            DENTAL COVERAGE POLICY
            
            Covered Procedures:
            - Routine dental checkups (2 per year)
            - Dental fillings and extractions
            - Root canal treatment (up to $2,000)
            - Dental surgery for medical necessity
            - Wisdom tooth extraction
            
            Coverage Limits:
            - Annual maximum: $3,000
            - Waiting period: 3 months for major procedures
            - Co-pay: 20% for specialist consultations
            
            Age Eligibility:
            - Minimum age: 18 years
            - Maximum age: 75 years
            - Children under 18 covered under family plan
            
            Network Dentists:
            - All registered dental clinics
            - Emergency dental care covered
            - Coverage in all major cities
            """,
            "source": "dental_policy.txt"
        },
        {
            "content": """
            CARDIAC CARE COVERAGE POLICY
            
            Heart Surgery Coverage:
            - Bypass surgery: Up to $100,000
            - Angioplasty: Up to $75,000
            - Heart valve replacement: Up to $150,000
            - Emergency cardiac procedures: Full coverage
            - Cardiac rehabilitation: Up to $5,000
            
            Eligibility Requirements:
            - Age 18-70 for major cardiac procedures
            - Pre-existing heart conditions: 24-month waiting period
            - Second opinion required for surgeries over $50,000
            - Annual health checkup mandatory
            
            Waiting Periods:
            - 12 months for major cardiac procedures
            - 6 months for diagnostic procedures
            - Emergency procedures: No waiting period
            
            Network Cardiac Centers:
            - All accredited cardiac hospitals
            - Emergency air ambulance coverage
            - Coverage in all major cities including Pune, Mumbai, Delhi
            """,
            "source": "cardiac_policy.txt"
        },
        {
            "content": """
            GENERAL SURGERY COVERAGE POLICY
            
            Covered Procedures:
            - General surgery procedures: Up to $25,000
            - Appendectomy, hernia repair, gallbladder surgery
            - Minor surgical procedures: Up to $10,000
            - Emergency surgeries: Full coverage
            
            Age and Policy Requirements:
            - Minimum age: 18 years
            - Maximum age: 80 years
            - Policy must be active for 6 months minimum
            - Pre-authorization for non-emergency procedures
            
            Network Hospitals:
            - All major hospitals and medical centers
            - Coverage in all cities including Pune, Mumbai, Delhi
            - Emergency procedures covered immediately
            
            Exclusions:
            - Cosmetic procedures not covered
            - Experimental treatments excluded
            - Pre-existing conditions have waiting period
            """,
            "source": "general_surgery_policy.txt"
        }
    ]
    return sample_docs

def load_single_document(file_path: str):
    """Load a single document from file"""
    path_obj = Path(file_path)
    
    if path_obj.suffix.lower() == '.pdf':
        content = _extract_pdf_content(file_path)
    elif path_obj.suffix.lower() in ['.docx', '.doc']:
        content = _extract_docx_content(file_path)
    elif path_obj.suffix.lower() == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    else:
        raise ValueError(f"Unsupported file format: {file_path}")
    
    return Document(page_content=content, metadata={'source': file_path})

def _extract_pdf_content(path: str) -> str:
    """Extract text from PDF using PyMuPDF"""
    if fitz is None:
        raise ImportError("PyMuPDF not available. Install with: pip install PyMuPDF")
        
    try:
        doc = fitz.open(path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        raise Exception(f"Failed to extract PDF content: {e}")

def _extract_docx_content(path: str) -> str:
    """Extract text from DOCX"""
    if DocxDocument is None:
        raise ImportError("python-docx not available. Install with: pip install python-docx")
        
    try:
        doc = DocxDocument(path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise Exception(f"Failed to extract DOCX content: {e}")

def print_formatted_result(result: Dict[str, Any]):
    """Print results in a conversational format"""
    
    # Generate conversational response
    response = generate_conversational_response(result)
    print(f"\nRESPONSE: {response}")
    
    # Show confidence and key details
    confidence = result['confidence']
    decision = result['decision']
    
    if confidence < 0.5:
        confidence_note = " (Low confidence - manual review recommended)"
    elif confidence < 0.8:
        confidence_note = " (Moderate confidence)"
    else:
        confidence_note = " (High confidence)"
    
    print(f"Confidence: {confidence:.1%}{confidence_note}")
    
    # Show amount if available
    if result.get('amount'):
        print(f"Coverage Amount: ${result['amount']:,.2f}")
    
    # Show key reasoning if confidence is low
    if confidence < 0.7:
        print(f"\nKey Factors:")
        analysis = result["query_analysis"]
        if analysis.get("procedure"):
            print(f"   • Procedure: {analysis['procedure']}")
        if analysis.get("policy_age_months") and analysis['policy_age_months'] < 12:
            print(f"   • Policy Age: {analysis['policy_age_months']} months (may affect coverage)")
        if analysis.get("age") and analysis['age'] > 65:
            print(f"   • Age: {analysis['age']} years (senior citizen considerations)")
    
    print("-" * 60)

def generate_conversational_response(result: Dict[str, Any]) -> str:
    """Generate a conversational response based on the decision"""
    decision = result['decision']
    analysis = result['query_analysis']
    procedure = analysis.get('procedure', 'the procedure')
    amount = result.get('amount')
    confidence = result.get('confidence', 0)
    
    # Build response based on decision and confidence
    if decision == "approved":
        response = f"Yes, {procedure} is covered under the policy."
        if amount:
            response += f" Coverage amount: ${amount:,.2f}."
        if confidence < 0.8:
            response += " (Manual review recommended for final confirmation)"
        return response
    
    elif decision == "rejected":
        response = f"No, {procedure} is not covered under the policy."
        if confidence < 0.8:
            response += " (Manual review recommended for final confirmation)"
        return response
    
    elif decision == "requires_review":
        response = f"This {procedure} requires manual review."
        if analysis.get('policy_age_months') and analysis['policy_age_months'] < 6:
            response += f" Policy is only {analysis['policy_age_months']} months old and may have waiting period restrictions."
        elif analysis.get('age') and analysis['age'] > 65:
            response += f" Age {analysis['age']} may have senior citizen considerations."
        else:
            response += " Please contact customer service for detailed assessment."
        return response
    
    else:
        return "Unable to determine coverage. Please contact customer service for assistance."

def main(document_path: str = None):
    """Example usage of the Enhanced RAG System"""
    
    # Initialize system
    rag_system = EnhancedRAGSystem()
    
    if document_path:
        # Load user's document using enhanced processing
        print(f"Loading document: {document_path}")
        try:
            # Use the enhanced document loading
            rag_system.retriever.load_documents([document_path])
            print(f"Document loaded successfully!")
        except Exception as e:
            print(f"Error loading document: {e}")
            return
    else:
        # Create and load sample documents
        print("Loading sample documents...")
        sample_docs = create_sample_documents()
        
        # Convert sample documents to Document objects and load them
        if Document is not None:
            documents = []
            for doc in sample_docs:
                documents.append(Document(page_content=doc["content"], metadata={'source': doc["source"]}))
            
            # Load sample documents
            rag_system.retriever.documents = documents
            rag_system.retriever.chunks = documents  # For sample docs, use full documents
    
    # Interactive query testing
    print("\n" + "="*60)
    print("RAG SYSTEM READY FOR TESTING")
    print("="*60)
    print("Enter your queries (type 'quit' to exit):")
    
    while True:
        try:
            query = input("\nEnter your query: ").strip()
            if query.lower() in ['quit', 'exit', 'q']:
                break
            if not query:
                continue
                
            print(f"\nProcessing: {query}")
            print("-" * 50)
            
            result = rag_system.process_query(query, debug=True)
            print_formatted_result(result)
            
        except KeyboardInterrupt:
            print("\nExiting...")
            break
        except Exception as e:
            print(f"Error: {e}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        # Use document path from command line argument
        main(sys.argv[1])
    else:
        # Use sample documents
        main()